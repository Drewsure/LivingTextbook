from fastapi import FastAPI, APIRouter, HTTPException, Depends, status, UploadFile, File, Form
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import bcrypt
import jwt
from emergentintegrations.llm.chat import LlmChat, UserMessage, FileContentWithMimeType
import base64
import io
from PyPDF2 import PdfReader
import docx
import qrcode
import secrets
import string
import json
from fpdf import FPDF
from fastapi.responses import Response
import fitz  # PyMuPDF for image extraction
from PIL import Image

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

mongo_url = os.environ.get('MONGO_URL')
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ.get('DB_NAME')]

app = FastAPI()
api_router = APIRouter(prefix="/api")
security = HTTPBearer()

JWT_SECRET = os.environ.get('JWT_SECRET', 'secret')
JWT_ALGORITHM = 'HS256'
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY')

# Models
class UserRegister(BaseModel):
    email: EmailStr
    password: str
    name: str
    role: str  # 'teacher' or 'student'

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: str
    name: str
    role: str
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class TokenResponse(BaseModel):
    token: str
    user: User

class WorksheetAnalysis(BaseModel):
    keywords: List[str]
    theme: str
    grade_level: str
    learning_objectives: List[str]
    suggested_games: List[str]
    image_keywords: Optional[List[str]] = []  # Keywords for image search

class Worksheet(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    teacher_id: str
    title: str
    content: str
    file_type: str
    analysis: Optional[Dict] = None
    instruction_language: str = "ja"  # Default to Japanese
    extracted_images: Optional[List[str]] = []  # Base64 encoded images from worksheet
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class GameCreate(BaseModel):
    worksheet_id: str
    title: str
    game_type: str
    content: Dict
    grade_level: Optional[str] = None
    theme: Optional[str] = None

class Game(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    worksheet_id: str
    teacher_id: str
    title: str
    game_type: str
    content: Dict
    grade_level: Optional[str] = None
    theme: Optional[str] = None
    share_code: str = Field(default_factory=lambda: ''.join(secrets.choice(string.ascii_uppercase + string.digits) for _ in range(6)))
    qr_code: Optional[str] = None
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())
    plays: int = 0
    is_favorite: bool = False

class Leaderboard(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    student_name: str
    score: int
    game_id: str
    game_title: str
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class GameSession(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    game_id: str
    student_id: Optional[str] = None
    student_name: Optional[str] = None
    score: int
    completed: bool
    time_spent: int  # seconds
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class StudentProgress(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    student_id: str
    student_name: str
    total_games_played: int = 0
    total_score: int = 0
    games_completed: int = 0
    average_score: float = 0.0
    best_score: int = 0
    favorite_game_type: Optional[str] = None
    last_played: Optional[str] = None
    badges: List[str] = []
    streak_days: int = 0
    created_at: str = Field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

class TeacherStats(BaseModel):
    active_students: int
    worksheets_created: int
    games_played: int
    average_engagement: int

# Auth utilities
def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def create_token(user_id: str, email: str, role: str) -> str:
    payload = {
        'user_id': user_id,
        'email': email,
        'role': role,
        'exp': datetime.now(timezone.utc) + timedelta(days=30)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        token = credentials.credentials
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user = await db.users.find_one({'id': payload['user_id']}, {'_id': 0})
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        return User(**user)
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except Exception as e:
        raise HTTPException(status_code=401, detail="Invalid token")

# Auth endpoints
@api_router.post("/auth/register", response_model=TokenResponse)
async def register(user_data: UserRegister):
    existing = await db.users.find_one({'email': user_data.email}, {'_id': 0})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user_dict = {
        'id': str(uuid.uuid4()),
        'email': user_data.email,
        'password': hash_password(user_data.password),
        'name': user_data.name,
        'role': user_data.role,
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    
    await db.users.insert_one(user_dict)
    
    user = User(**{k: v for k, v in user_dict.items() if k != 'password'})
    token = create_token(user.id, user.email, user.role)
    
    return TokenResponse(token=token, user=user)

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(credentials: UserLogin):
    user = await db.users.find_one({'email': credentials.email}, {'_id': 0})
    if not user or not verify_password(credentials.password, user['password']):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    user_obj = User(**{k: v for k, v in user.items() if k != 'password'})
    token = create_token(user_obj.id, user_obj.email, user_obj.role)
    
    return TokenResponse(token=token, user=user_obj)

@api_router.get("/auth/me", response_model=User)
async def get_me(current_user: User = Depends(get_current_user)):
    return current_user

# Image extraction helper functions
def extract_images_from_pdf(file_content: bytes, max_images: int = 10) -> List[str]:
    """Extract images from PDF and return as base64 strings"""
    images = []
    try:
        pdf_document = fitz.open(stream=file_content, filetype="pdf")
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                if len(images) >= max_images:
                    break
                    
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                
                # Convert to base64
                img_base64 = base64.b64encode(image_bytes).decode('utf-8')
                ext = base_image["ext"]
                data_url = f"data:image/{ext};base64,{img_base64}"
                images.append(data_url)
                
            if len(images) >= max_images:
                break
        pdf_document.close()
    except Exception as e:
        logging.error(f"Error extracting images from PDF: {e}")
    return images

def extract_images_from_docx(file_content: bytes, max_images: int = 10) -> List[str]:
    """Extract images from DOCX and return as base64 strings"""
    images = []
    try:
        doc = docx.Document(io.BytesIO(file_content))
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    # Determine image type
                    if rel.target_ref.endswith('.png'):
                        mime = 'image/png'
                    elif rel.target_ref.endswith('.jpg') or rel.target_ref.endswith('.jpeg'):
                        mime = 'image/jpeg'
                    elif rel.target_ref.endswith('.gif'):
                        mime = 'image/gif'
                    else:
                        mime = 'image/png'
                    
                    img_base64 = base64.b64encode(image_data).decode('utf-8')
                    data_url = f"data:{mime};base64,{img_base64}"
                    images.append(data_url)
                    
                    if len(images) >= max_images:
                        break
                except Exception as e:
                    logging.error(f"Error extracting image from DOCX: {e}")
    except Exception as e:
        logging.error(f"Error processing DOCX for images: {e}")
    return images

def process_uploaded_image(file_content: bytes, content_type: str, max_images: int = 10) -> List[str]:
    """Process an uploaded image file and return as base64 string"""
    images = []
    try:
        img_base64 = base64.b64encode(file_content).decode('utf-8')
        data_url = f"data:{content_type};base64,{img_base64}"
        images.append(data_url)
    except Exception as e:
        logging.error(f"Error processing uploaded image: {e}")
    return images

# Worksheet endpoints
async def analyze_worksheet_with_ai(content: str, file_type: str) -> WorksheetAnalysis:
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=str(uuid.uuid4()),
            system_message="You are an expert ESL/EFL education analyst. Your PRIMARY task is to identify the EXACT learning focus of the worksheet by carefully reading the context, not just counting words."
        ).with_model("gemini", "gemini-2.5-pro")
        
        prompt = f"""CRITICAL: Read this worksheet CAREFULLY and identify its PRIMARY learning focus.

ANALYSIS RULES:
1. If the worksheet says "Personal Pronouns" or "Subject Pronouns" → theme is PRONOUNS
2. If it lists I, you, he, she, it, we, they with examples → theme is PRONOUNS
3. If it shows colors (red, blue, green) with examples → theme is COLORS
4. If it shows animals with pictures/descriptions → theme is ANIMALS
5. If it shows family members (mother, father, sister) → theme is FAMILY
6. If it shows food items → theme is FOOD
7. Look at the TITLE and INTRODUCTION to understand the focus

WORKSHEET CONTENT:
{content[:4000]}

Provide JSON response with these fields:
- keywords: array of 10-15 ACTUAL vocabulary words from the worksheet (the words students should learn)
- theme: The PRIMARY learning focus (e.g., "Pronouns", "Colors", "Animals", "Family", "Food", "Verbs", "Adjectives")
- grade_level: appropriate grade (Kindergarten, Grade 1, Grade 2, Grade 3, Grade 4, Grade 5)
- learning_objectives: array of 4-6 specific, actionable learning objectives related to the theme
- suggested_games: array of 6-8 game types: Word Match, Quiz, Flashcards, Spelling Practice, Word Search, Matching Pairs, Fill in the Blank, True or False, Word Climber
- image_keywords: array of 10-15 simple nouns for finding relevant images that match the THEME (e.g., for Pronouns: "person", "student", "teacher", "friends", "family", "children", "boy", "girl", "people", "classroom")

RESPOND ONLY WITH VALID JSON. NO EXPLANATIONS."""
        
        message = UserMessage(text=prompt)
        response = await chat.send_message(message)
        
        # Parse AI response
        response_text = response.strip()
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.startswith('```'):
            response_text = response_text[3:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
        
        analysis_data = json.loads(response_text.strip())
        
        return WorksheetAnalysis(**analysis_data)
    except Exception as e:
        logging.error(f"AI analysis error: {e}")
        # Fallback analysis
        return WorksheetAnalysis(
            keywords=["vocabulary", "practice", "learning"],
            theme="General",
            grade_level="Grade 2",
            learning_objectives=["Practice vocabulary", "Improve comprehension"],
            suggested_games=["Word Match", "Quiz", "Flashcards", "Spelling Practice", "Word Search", "Word Climber"],
            image_keywords=["book", "pencil", "child", "teacher", "classroom"]
        )

# AI Theme Monitoring - Log analysis for accuracy tracking
async def log_theme_analysis(worksheet_id: str, content_preview: str, detected_theme: str, keywords: List[str]):
    """Log AI theme analysis for monitoring accuracy"""
    log_entry = {
        'id': str(uuid.uuid4()),
        'worksheet_id': worksheet_id,
        'content_preview': content_preview[:200],  # First 200 chars
        'detected_theme': detected_theme,
        'keywords': keywords[:10],
        'timestamp': datetime.now(timezone.utc).isoformat()
    }
    await db.theme_analysis_logs.insert_one(log_entry)
    logging.info(f"Theme analysis logged: {detected_theme} for worksheet {worksheet_id}")

# Endpoint to get theme analysis stats for monitoring
@api_router.get("/admin/theme-analysis-stats")
async def get_theme_analysis_stats(current_user: User = Depends(get_current_user)):
    """Get statistics on AI theme analysis for monitoring accuracy"""
    if current_user.role != 'teacher':
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Get all theme analysis logs
    logs = await db.theme_analysis_logs.find({}, {'_id': 0}).sort('timestamp', -1).limit(100).to_list(100)
    
    # Calculate theme distribution
    theme_counts = {}
    for log in logs:
        theme = log.get('detected_theme', 'Unknown')
        theme_counts[theme] = theme_counts.get(theme, 0) + 1
    
    # Sort by count
    theme_distribution = [
        {'theme': k, 'count': v}
        for k, v in sorted(theme_counts.items(), key=lambda x: x[1], reverse=True)
    ]
    
    return {
        'total_analyses': len(logs),
        'theme_distribution': theme_distribution,
        'recent_analyses': logs[:20]
    }

@api_router.post("/worksheets/upload")
async def upload_worksheet(
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    instruction_language: str = Form("ja"),
    current_user: User = Depends(get_current_user)
):
    if current_user.role != 'teacher':
        raise HTTPException(status_code=403, detail="Only teachers can upload worksheets")
    
    file_content = await file.read()
    file_type = file.content_type
    
    # Extract text and images based on file type
    text_content = ""
    extracted_images = []
    
    if file_type == "application/pdf":
        pdf_reader = PdfReader(io.BytesIO(file_content))
        text_content = " ".join([page.extract_text() for page in pdf_reader.pages])
        # Extract images from PDF
        extracted_images = extract_images_from_pdf(file_content)
        
    elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        doc = docx.Document(io.BytesIO(file_content))
        text_content = " ".join([para.text for para in doc.paragraphs])
        # Extract images from DOCX
        extracted_images = extract_images_from_docx(file_content)
        
    elif file_type.startswith("text/"):
        text_content = file_content.decode('utf-8')
        
    elif file_type.startswith("image/"):
        # For images, use Gemini's vision capabilities and store the image
        extracted_images = process_uploaded_image(file_content, file_type)
        try:
            temp_path = f"/tmp/{uuid.uuid4()}_{file.filename}"
            with open(temp_path, 'wb') as f:
                f.write(file_content)
            
            chat = LlmChat(
                api_key=EMERGENT_LLM_KEY,
                session_id=str(uuid.uuid4()),
                system_message="You are an OCR and education content extractor."
            ).with_model("gemini", "gemini-2.5-pro")
            
            file_obj = FileContentWithMimeType(
                file_path=temp_path,
                mime_type=file_type
            )
            
            message = UserMessage(
                text="Extract all text content from this worksheet image. Provide the raw text.",
                file_contents=[file_obj]
            )
            
            text_content = await chat.send_message(message)
            
            os.remove(temp_path)
        except Exception as e:
            logging.error(f"Image OCR error: {e}")
            text_content = "[Image content - manual review needed]"
    
    # Analyze content
    analysis = await analyze_worksheet_with_ai(text_content, file_type)
    
    worksheet_id = str(uuid.uuid4())
    
    # Log theme analysis for monitoring
    await log_theme_analysis(
        worksheet_id=worksheet_id,
        content_preview=text_content,
        detected_theme=analysis.theme,
        keywords=analysis.keywords
    )
    
    worksheet_dict = {
        'id': worksheet_id,
        'teacher_id': current_user.id,
        'title': title or file.filename,
        'content': text_content[:5000],  # Store first 5000 chars
        'file_type': file_type,
        'analysis': analysis.model_dump(),
        'instruction_language': instruction_language,
        'extracted_images': extracted_images,  # Store extracted images
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    
    await db.worksheets.insert_one(worksheet_dict)
    
    return Worksheet(**worksheet_dict)

@api_router.post("/worksheets/analyze-text")
async def analyze_text(
    text: str = Form(...),
    title: str = Form(...),
    instruction_language: str = Form("ja"),
    current_user: User = Depends(get_current_user)
):
    if current_user.role != 'teacher':
        raise HTTPException(status_code=403, detail="Only teachers can create worksheets")
    
    analysis = await analyze_worksheet_with_ai(text, "text/plain")
    
    worksheet_id = str(uuid.uuid4())
    
    # Log theme analysis for monitoring
    await log_theme_analysis(
        worksheet_id=worksheet_id,
        content_preview=text,
        detected_theme=analysis.theme,
        keywords=analysis.keywords
    )
    
    worksheet_dict = {
        'id': worksheet_id,
        'teacher_id': current_user.id,
        'title': title,
        'content': text[:5000],
        'file_type': 'text/plain',
        'analysis': analysis.model_dump(),
        'instruction_language': instruction_language,
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    
    await db.worksheets.insert_one(worksheet_dict)
    
    return Worksheet(**worksheet_dict)

@api_router.post("/worksheets/{worksheet_id}/reanalyze")
async def reanalyze_worksheet(worksheet_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != 'teacher':
        raise HTTPException(status_code=403, detail="Only teachers can reanalyze worksheets")
    
    worksheet = await db.worksheets.find_one({'id': worksheet_id, 'teacher_id': current_user.id}, {'_id': 0})
    if not worksheet:
        raise HTTPException(status_code=404, detail="Worksheet not found")
    
    # Re-analyze
    analysis = await analyze_worksheet_with_ai(worksheet['content'], worksheet['file_type'])
    
    # Update worksheet
    await db.worksheets.update_one(
        {'id': worksheet_id},
        {'$set': {'analysis': analysis.model_dump()}}
    )
    
    worksheet['analysis'] = analysis.model_dump()
    return Worksheet(**worksheet)

@api_router.get("/worksheets", response_model=List[Worksheet])
async def get_worksheets(current_user: User = Depends(get_current_user)):
    if current_user.role != 'teacher':
        raise HTTPException(status_code=403, detail="Only teachers can view worksheets")
    
    worksheets = await db.worksheets.find({'teacher_id': current_user.id}, {'_id': 0}).sort('created_at', -1).to_list(100)
    return [Worksheet(**w) for w in worksheets]

@api_router.get("/worksheets/{worksheet_id}", response_model=Worksheet)
async def get_worksheet(worksheet_id: str, current_user: User = Depends(get_current_user)):
    worksheet = await db.worksheets.find_one({'id': worksheet_id}, {'_id': 0})
    if not worksheet:
        raise HTTPException(status_code=404, detail="Worksheet not found")
    
    if worksheet['teacher_id'] != current_user.id and current_user.role != 'student':
        raise HTTPException(status_code=403, detail="Access denied")
    
    return Worksheet(**worksheet)

# Endpoint to get extracted images for a worksheet (public for games)
@api_router.get("/worksheets/{worksheet_id}/images")
async def get_worksheet_images(worksheet_id: str):
    """Get extracted images from a worksheet for use in games"""
    worksheet = await db.worksheets.find_one({'id': worksheet_id}, {'_id': 0})
    if not worksheet:
        raise HTTPException(status_code=404, detail="Worksheet not found")
    
    return {
        "worksheet_id": worksheet_id,
        "images": worksheet.get('extracted_images', []),
        "image_count": len(worksheet.get('extracted_images', []))
    }

# Game endpoints
@api_router.post("/games/create", response_model=Game)
async def create_game(game_data: GameCreate, current_user: User = Depends(get_current_user)):
    if current_user.role != 'teacher':
        raise HTTPException(status_code=403, detail="Only teachers can create games")
    
    worksheet = await db.worksheets.find_one({'id': game_data.worksheet_id}, {'_id': 0})
    if not worksheet or worksheet['teacher_id'] != current_user.id:
        raise HTTPException(status_code=404, detail="Worksheet not found")
    
    share_code = ''.join(secrets.choice(string.ascii_uppercase + string.digits) for _ in range(6))
    
    # Generate QR code
    qr = qrcode.QRCode(version=1, box_size=10, border=5)
    qr.add_data(f"https://ai-quiz-builder-1.preview.emergentagent.com/play/{share_code}")
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    
    buffer = io.BytesIO()
    img.save(buffer, format='PNG')
    qr_code_base64 = base64.b64encode(buffer.getvalue()).decode()
    
    # Include worksheet images in game content
    game_content = game_data.content.copy()
    game_content['worksheetImages'] = worksheet.get('extracted_images', [])
    
    game_dict = {
        'id': str(uuid.uuid4()),
        'worksheet_id': game_data.worksheet_id,
        'teacher_id': current_user.id,
        'title': game_data.title,
        'game_type': game_data.game_type,
        'content': game_content,
        'grade_level': game_data.grade_level,
        'theme': game_data.theme,
        'share_code': share_code,
        'qr_code': qr_code_base64,
        'created_at': datetime.now(timezone.utc).isoformat(),
        'plays': 0
    }
    
    await db.games.insert_one(game_dict)
    
    return Game(**game_dict)

@api_router.get("/games/by-code/{share_code}", response_model=Game)
async def get_game_by_code(share_code: str):
    game = await db.games.find_one({'share_code': share_code.upper()}, {'_id': 0})
    if not game:
        raise HTTPException(status_code=404, detail="Game not found")
    
    # Increment play count
    await db.games.update_one({'share_code': share_code.upper()}, {'$inc': {'plays': 1}})
    
    return Game(**game)

@api_router.get("/games/worksheet/{worksheet_id}", response_model=List[Game])
async def get_games_by_worksheet(worksheet_id: str, current_user: User = Depends(get_current_user)):
    games = await db.games.find({'worksheet_id': worksheet_id}, {'_id': 0}).to_list(100)
    return [Game(**g) for g in games]

@api_router.get("/games/teacher", response_model=List[Game])
async def get_teacher_games(current_user: User = Depends(get_current_user)):
    if current_user.role != 'teacher':
        raise HTTPException(status_code=403, detail="Only teachers can view their games")
    
    # Sort by is_favorite (desc) then created_at (desc) so favorites appear first
    games = await db.games.find({'teacher_id': current_user.id}, {'_id': 0}).sort([('is_favorite', -1), ('created_at', -1)]).to_list(100)
    return [Game(**g) for g in games]

# Toggle favorite status
@api_router.post("/games/{game_id}/favorite")
async def toggle_favorite(game_id: str, current_user: User = Depends(get_current_user)):
    if current_user.role != 'teacher':
        raise HTTPException(status_code=403, detail="Only teachers can favorite games")
    
    game = await db.games.find_one({'id': game_id, 'teacher_id': current_user.id}, {'_id': 0})
    if not game:
        raise HTTPException(status_code=404, detail="Game not found")
    
    new_status = not game.get('is_favorite', False)
    await db.games.update_one(
        {'id': game_id},
        {'$set': {'is_favorite': new_status}}
    )
    
    game['is_favorite'] = new_status
    return Game(**game)

# Get favorite games only
@api_router.get("/games/favorites", response_model=List[Game])
async def get_favorite_games(current_user: User = Depends(get_current_user)):
    if current_user.role != 'teacher':
        raise HTTPException(status_code=403, detail="Only teachers can view favorites")
    
    games = await db.games.find(
        {'teacher_id': current_user.id, 'is_favorite': True}, 
        {'_id': 0}
    ).sort('plays', -1).to_list(100)
    return [Game(**g) for g in games]

# Game session endpoints
@api_router.post("/game-sessions", response_model=GameSession)
async def create_game_session(session_data: GameSession):
    session_dict = session_data.model_dump()
    await db.game_sessions.insert_one(session_dict)
    return session_data

@api_router.get("/game-sessions/game/{game_id}", response_model=List[GameSession])
async def get_game_sessions(game_id: str, current_user: User = Depends(get_current_user)):
    sessions = await db.game_sessions.find({'game_id': game_id}, {'_id': 0}).sort('created_at', -1).to_list(100)
    return [GameSession(**s) for s in sessions]

# Teacher stats
@api_router.get("/teacher/stats", response_model=TeacherStats)
async def get_teacher_stats(current_user: User = Depends(get_current_user)):
    if current_user.role != 'teacher':
        raise HTTPException(status_code=403, detail="Only teachers can view stats")
    
    worksheets_count = await db.worksheets.count_documents({'teacher_id': current_user.id})
    games = await db.games.find({'teacher_id': current_user.id}, {'_id': 0}).to_list(1000)
    total_plays = sum(g.get('plays', 0) for g in games)
    
    sessions = await db.game_sessions.find({}, {'_id': 0}).to_list(10000)
    unique_students = len(set(s.get('student_id') for s in sessions if s.get('student_id')))
    
    avg_engagement = 0
    if sessions:
        completed = [s for s in sessions if s.get('completed')]
        if completed:
            avg_engagement = int(sum(s.get('score', 0) for s in completed) / len(completed))
    
    return TeacherStats(
        active_students=unique_students,
        worksheets_created=worksheets_count,
        games_played=total_plays,
        average_engagement=avg_engagement
    )

# Analytics endpoints
@api_router.get("/teacher/analytics")
async def get_teacher_analytics(current_user: User = Depends(get_current_user)):
    """Get comprehensive analytics for teacher's games"""
    if current_user.role != 'teacher':
        raise HTTPException(status_code=403, detail="Only teachers can view analytics")
    
    # Get all teacher's games
    games = await db.games.find({'teacher_id': current_user.id}, {'_id': 0}).to_list(1000)
    game_ids = [g['id'] for g in games]
    
    # Get all sessions for these games
    sessions = await db.game_sessions.find({'game_id': {'$in': game_ids}}, {'_id': 0}).to_list(10000)
    
    # Get leaderboard entries for these games
    leaderboard_entries = await db.leaderboard.find({'game_id': {'$in': game_ids}}, {'_id': 0}).to_list(10000)
    
    # Calculate per-game metrics
    game_metrics = []
    for game in games:
        game_sessions = [s for s in sessions if s.get('game_id') == game['id']]
        game_leaderboard = [l for l in leaderboard_entries if l.get('game_id') == game['id']]
        
        completed_sessions = [s for s in game_sessions if s.get('completed')]
        total_plays = game.get('plays', 0)
        
        avg_score = 0
        if game_leaderboard:
            avg_score = sum(l.get('score', 0) for l in game_leaderboard) / len(game_leaderboard)
        elif completed_sessions:
            avg_score = sum(s.get('score', 0) for s in completed_sessions) / len(completed_sessions)
        
        avg_time = 0
        if completed_sessions:
            avg_time = sum(s.get('time_spent', 0) for s in completed_sessions) / len(completed_sessions)
        
        completion_rate = 0
        if game_sessions:
            completion_rate = (len(completed_sessions) / len(game_sessions)) * 100
        
        game_metrics.append({
            'id': game['id'],
            'title': game['title'],
            'game_type': game['game_type'],
            'theme': game.get('theme', 'General'),
            'total_plays': total_plays,
            'unique_players': len(set(s.get('student_id') for s in game_sessions if s.get('student_id'))),
            'avg_score': round(avg_score, 1),
            'avg_time_seconds': round(avg_time, 0),
            'completion_rate': round(completion_rate, 1),
            'is_favorite': game.get('is_favorite', False),
            'created_at': game.get('created_at', '')
        })
    
    # Sort by total plays for ranking
    game_metrics.sort(key=lambda x: x['total_plays'], reverse=True)
    
    # Calculate overall metrics
    total_plays = sum(g['total_plays'] for g in game_metrics)
    total_completed = len([s for s in sessions if s.get('completed')])
    overall_avg_score = 0
    if leaderboard_entries:
        overall_avg_score = sum(l.get('score', 0) for l in leaderboard_entries) / len(leaderboard_entries)
    
    # Game type distribution
    game_type_stats = {}
    for game in games:
        gt = game['game_type']
        if gt not in game_type_stats:
            game_type_stats[gt] = {'count': 0, 'plays': 0}
        game_type_stats[gt]['count'] += 1
        game_type_stats[gt]['plays'] += game.get('plays', 0)
    
    game_type_distribution = [
        {'name': k, 'games': v['count'], 'plays': v['plays']}
        for k, v in game_type_stats.items()
    ]
    game_type_distribution.sort(key=lambda x: x['plays'], reverse=True)
    
    # Theme distribution  
    theme_stats = {}
    for game in games:
        theme = game.get('theme', 'General')
        if theme not in theme_stats:
            theme_stats[theme] = {'count': 0, 'plays': 0}
        theme_stats[theme]['count'] += 1
        theme_stats[theme]['plays'] += game.get('plays', 0)
    
    theme_distribution = [
        {'name': k, 'games': v['count'], 'plays': v['plays']}
        for k, v in theme_stats.items()
    ]
    theme_distribution.sort(key=lambda x: x['plays'], reverse=True)
    
    # Top performers
    top_by_plays = sorted(game_metrics, key=lambda x: x['total_plays'], reverse=True)[:5]
    top_by_score = sorted([g for g in game_metrics if g['avg_score'] > 0], key=lambda x: x['avg_score'], reverse=True)[:5]
    top_by_completion = sorted([g for g in game_metrics if g['completion_rate'] > 0], key=lambda x: x['completion_rate'], reverse=True)[:5]
    
    return {
        'summary': {
            'total_games': len(games),
            'total_plays': total_plays,
            'total_completed': total_completed,
            'overall_avg_score': round(overall_avg_score, 1),
            'unique_players': len(set(s.get('student_id') for s in sessions if s.get('student_id')))
        },
        'game_metrics': game_metrics,
        'game_type_distribution': game_type_distribution,
        'theme_distribution': theme_distribution,
        'top_performers': {
            'by_plays': top_by_plays,
            'by_score': top_by_score,
            'by_completion': top_by_completion
        }
    }

# Leaderboard endpoints
@api_router.post("/leaderboard", response_model=Leaderboard)
async def add_to_leaderboard(entry: Leaderboard):
    entry_dict = entry.model_dump()
    await db.leaderboard.insert_one(entry_dict)
    return entry

@api_router.get("/leaderboard/top", response_model=List[Leaderboard])
async def get_top_leaderboard(limit: int = 15):
    entries = await db.leaderboard.find({}, {'_id': 0}).sort('score', -1).limit(limit).to_list(limit)
    return [Leaderboard(**e) for e in entries]

@api_router.get("/leaderboard/game/{game_id}", response_model=List[Leaderboard])
async def get_game_leaderboard(game_id: str, limit: int = 15):
    entries = await db.leaderboard.find({'game_id': game_id}, {'_id': 0}).sort('score', -1).limit(limit).to_list(limit)
    return [Leaderboard(**e) for e in entries]

# Student Progress Tracking endpoints
@api_router.post("/student/progress/update")
async def update_student_progress(
    student_name: str,
    game_id: str,
    score: int,
    completed: bool,
    time_spent: int
):
    """Update student progress after playing a game"""
    # Create a unique student ID based on name (for anonymous tracking)
    student_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, student_name.lower().strip()))
    
    # Get game details
    game = await db.games.find_one({'id': game_id}, {'_id': 0})
    
    # Record the session
    session_dict = {
        'id': str(uuid.uuid4()),
        'game_id': game_id,
        'student_id': student_id,
        'student_name': student_name,
        'score': score,
        'completed': completed,
        'time_spent': time_spent,
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    await db.game_sessions.insert_one(session_dict)
    
    # Update or create student progress
    existing_progress = await db.student_progress.find_one({'student_id': student_id}, {'_id': 0})
    
    today = datetime.now(timezone.utc).date().isoformat()
    
    if existing_progress:
        # Update existing progress
        total_games = existing_progress.get('total_games_played', 0) + 1
        total_score = existing_progress.get('total_score', 0) + score
        games_completed = existing_progress.get('games_completed', 0) + (1 if completed else 0)
        best_score = max(existing_progress.get('best_score', 0), score)
        avg_score = total_score / total_games if total_games > 0 else 0
        
        # Calculate streak
        last_played = existing_progress.get('last_played', '')
        streak = existing_progress.get('streak_days', 0)
        if last_played:
            last_date = datetime.fromisoformat(last_played.replace('Z', '+00:00')).date()
            today_date = datetime.now(timezone.utc).date()
            if (today_date - last_date).days == 1:
                streak += 1
            elif (today_date - last_date).days > 1:
                streak = 1
        else:
            streak = 1
        
        # Award badges
        badges = existing_progress.get('badges', [])
        if total_games >= 10 and 'starter' not in badges:
            badges.append('starter')
        if total_games >= 50 and 'dedicated' not in badges:
            badges.append('dedicated')
        if total_games >= 100 and 'master' not in badges:
            badges.append('master')
        if best_score >= 90 and 'perfectionist' not in badges:
            badges.append('perfectionist')
        if streak >= 7 and 'streak_week' not in badges:
            badges.append('streak_week')
        if streak >= 30 and 'streak_month' not in badges:
            badges.append('streak_month')
        
        await db.student_progress.update_one(
            {'student_id': student_id},
            {'$set': {
                'total_games_played': total_games,
                'total_score': total_score,
                'games_completed': games_completed,
                'average_score': round(avg_score, 1),
                'best_score': best_score,
                'favorite_game_type': game.get('game_type') if game else None,
                'last_played': datetime.now(timezone.utc).isoformat(),
                'badges': badges,
                'streak_days': streak
            }}
        )
    else:
        # Create new progress record
        badges = []
        if score >= 90:
            badges.append('perfectionist')
            
        progress_dict = {
            'id': str(uuid.uuid4()),
            'student_id': student_id,
            'student_name': student_name,
            'total_games_played': 1,
            'total_score': score,
            'games_completed': 1 if completed else 0,
            'average_score': float(score),
            'best_score': score,
            'favorite_game_type': game.get('game_type') if game else None,
            'last_played': datetime.now(timezone.utc).isoformat(),
            'badges': badges,
            'streak_days': 1,
            'created_at': datetime.now(timezone.utc).isoformat()
        }
        await db.student_progress.insert_one(progress_dict)
    
    # Get updated progress
    progress = await db.student_progress.find_one({'student_id': student_id}, {'_id': 0})
    return StudentProgress(**progress)

@api_router.get("/student/progress/{student_name}")
async def get_student_progress(student_name: str):
    """Get progress for a specific student"""
    student_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, student_name.lower().strip()))
    progress = await db.student_progress.find_one({'student_id': student_id}, {'_id': 0})
    
    if not progress:
        raise HTTPException(status_code=404, detail="Student not found")
    
    # Get recent sessions
    sessions = await db.game_sessions.find(
        {'student_id': student_id}, 
        {'_id': 0}
    ).sort('created_at', -1).limit(10).to_list(10)
    
    return {
        'progress': StudentProgress(**progress),
        'recent_sessions': sessions
    }

@api_router.get("/student/leaderboard")
async def get_student_leaderboard(limit: int = 20):
    """Get top students by total score"""
    students = await db.student_progress.find(
        {}, 
        {'_id': 0}
    ).sort('total_score', -1).limit(limit).to_list(limit)
    
    return [StudentProgress(**s) for s in students]

# Games Landing Page endpoint (for QR code access - public)
@api_router.get("/worksheets/landing/{worksheet_id}")
async def get_worksheet_landing_page(worksheet_id: str):
    """Public endpoint for students to access all games for a worksheet"""
    worksheet = await db.worksheets.find_one({'id': worksheet_id}, {'_id': 0})
    if not worksheet:
        raise HTTPException(status_code=404, detail="Worksheet not found")
    
    games = await db.games.find({'worksheet_id': worksheet_id}, {'_id': 0}).to_list(100)
    
    return {
        "worksheet": {
            "id": worksheet['id'],
            "title": worksheet['title'],
            "theme": worksheet.get('analysis', {}).get('theme', 'General'),
            "grade_level": worksheet.get('analysis', {}).get('grade_level', 'All Levels'),
            "instruction_language": worksheet.get('instruction_language', 'ja')
        },
        "games": [
            {
                "id": g['id'],
                "title": g['title'],
                "game_type": g['game_type'],
                "theme": g.get('theme'),
                "share_code": g['share_code'],
                "plays": g.get('plays', 0)
            }
            for g in games
        ]
    }

# Generate Printable QR Code PDF
@api_router.get("/worksheets/{worksheet_id}/print-qr")
async def generate_qr_pdf(worksheet_id: str, current_user: User = Depends(get_current_user)):
    """Generate a printable PDF with QR codes for all games"""
    if current_user.role != 'teacher':
        raise HTTPException(status_code=403, detail="Only teachers can generate QR PDFs")
    
    worksheet = await db.worksheets.find_one({'id': worksheet_id, 'teacher_id': current_user.id}, {'_id': 0})
    if not worksheet:
        raise HTTPException(status_code=404, detail="Worksheet not found")
    
    games = await db.games.find({'worksheet_id': worksheet_id}, {'_id': 0}).sort('created_at', 1).to_list(20)
    
    if not games:
        raise HTTPException(status_code=400, detail="No games found for this worksheet. Generate games first!")
    
    # Create PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Title
    pdf.set_font('Helvetica', 'B', 24)
    pdf.cell(0, 15, 'Ministar Game Studio', ln=True, align='C')
    
    pdf.set_font('Helvetica', 'B', 18)
    pdf.cell(0, 12, worksheet['title'], ln=True, align='C')
    
    theme = worksheet.get('analysis', {}).get('theme', 'General')
    grade = worksheet.get('analysis', {}).get('grade_level', 'All Levels')
    pdf.set_font('Helvetica', '', 12)
    pdf.cell(0, 8, f'Theme: {theme} | Level: {grade}', ln=True, align='C')
    
    pdf.ln(10)
    
    # Landing Page QR Code (all games)
    landing_url = f"https://ai-quiz-builder-1.preview.emergentagent.com/games/{worksheet_id}"
    
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, 'Scan to Play All Games', ln=True, align='C')
    
    # Generate main QR code
    main_qr = qrcode.QRCode(version=2, box_size=8, border=2)
    main_qr.add_data(landing_url)
    main_qr.make(fit=True)
    main_img = main_qr.make_image(fill_color="black", back_color="white")
    
    # Save temporarily
    main_qr_path = f"/tmp/main_qr_{worksheet_id}.png"
    main_img.save(main_qr_path)
    
    # Center the main QR code
    pdf.image(main_qr_path, x=70, w=70)
    pdf.ln(5)
    
    pdf.set_font('Helvetica', '', 10)
    pdf.cell(0, 8, landing_url, ln=True, align='C')
    
    pdf.ln(15)
    
    # Individual Game QR Codes (up to 6)
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, 'Individual Game Codes', ln=True, align='C')
    pdf.ln(5)
    
    # Display games in a grid (2 columns, up to 3 rows)
    games_to_show = games[:6]
    qr_paths = []
    
    for i, game in enumerate(games_to_show):
        game_url = f"https://ai-quiz-builder-1.preview.emergentagent.com/play/{game['share_code']}"
        
        qr = qrcode.QRCode(version=1, box_size=6, border=2)
        qr.add_data(game_url)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        
        qr_path = f"/tmp/game_qr_{game['id']}.png"
        img.save(qr_path)
        qr_paths.append(qr_path)
        
        # Calculate position (2 columns)
        col = i % 2
        if col == 0 and i > 0:
            pdf.ln(65)
        
        x_pos = 25 if col == 0 else 115
        
        # Game title
        pdf.set_xy(x_pos, pdf.get_y())
        pdf.set_font('Helvetica', 'B', 10)
        title = game['title'][:25] + '...' if len(game['title']) > 25 else game['title']
        pdf.cell(75, 6, title, ln=False, align='C')
        
        if col == 1:
            pdf.ln(6)
        
        # QR code
        pdf.image(qr_path, x=x_pos + 12, w=50)
        
        # Share code
        if col == 0:
            pdf.set_xy(x_pos, pdf.get_y() + 52)
        else:
            pdf.set_xy(x_pos, pdf.get_y())
        pdf.set_font('Helvetica', '', 9)
        pdf.cell(75, 5, f'Code: {game["share_code"]}', ln=False, align='C')
    
    # Add footer
    pdf.ln(20)
    pdf.set_font('Helvetica', 'I', 8)
    pdf.cell(0, 5, 'Generated by Ministar Game Studio - Making Learning Fun!', ln=True, align='C')
    
    # Clean up temp files
    try:
        os.remove(main_qr_path)
        for qr_path in qr_paths:
            os.remove(qr_path)
    except:
        pass
    
    # Return PDF
    pdf_bytes = pdf.output()
    
    return Response(
        content=bytes(pdf_bytes),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="ministar_games_{worksheet_id[:8]}.pdf"'
        }
    )

app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()